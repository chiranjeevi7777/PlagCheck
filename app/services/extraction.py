"""
Document text extraction service.

Supports PDF (PyMuPDF primary, pdfplumber fallback) and DOCX.
Strips headers, footers, page numbers, and normalises Unicode.
"""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path

import docx
import fitz          # PyMuPDF
import pdfplumber

from app.core.logging import get_logger

logger = get_logger(__name__)


class TextExtractor:
    """Extract and sanitise text from PDF and DOCX files."""

    @staticmethod
    def detect_and_extract(file_path: Path) -> str:
        """Detect file type, extract text, and return cleaned string."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            logger.info(f"Extracting PDF: {file_path.name}")
            raw = TextExtractor._extract_pdf(file_path)
        elif suffix in {".docx", ".doc"}:
            logger.info(f"Extracting DOCX: {file_path.name}")
            raw = TextExtractor._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
        return TextExtractor._clean(raw)

    # ── private ───────────────────────────────────────────────────────────────

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        try:
            doc = fitz.open(str(path))
            pages: list[str] = []
            for page in doc:
                rect = page.rect
                top_margin = rect.height * 0.08
                bottom_margin = rect.height * 0.92
                parts: list[str] = []
                for b in page.get_text("blocks"):
                    x0, y0, x1, y1, text, _, block_type = b
                    if block_type != 0 or not text.strip():
                        continue
                    if y0 < top_margin or y1 > bottom_margin:
                        t = text.strip()
                        if len(t) < 15 or re.match(r"^\d+$", t):
                            continue
                    parts.append(text.strip())
                pages.append("\n".join(parts))
            combined = "\n\n".join(pages)
            if combined.strip():
                return combined
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {path.name}, falling back: {e}")

        # pdfplumber fallback
        try:
            with pdfplumber.open(str(path)) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    lines = [
                        l.strip() for l in text.split("\n")
                        if l.strip() and not re.match(r"^\d+$", l.strip())
                    ]
                    pages.append("\n".join(lines))
            return "\n\n".join(pages)
        except Exception as e:
            raise RuntimeError(f"Could not extract text from PDF: {e}") from e

    @staticmethod
    def _extract_docx(path: Path) -> str:
        try:
            doc = docx.Document(str(path))
            parts: list[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)
        except Exception as e:
            raise RuntimeError(f"Could not extract text from DOCX: {e}") from e

    @staticmethod
    def _clean(text: str) -> str:
        if not text:
            return ""
        text = unicodedata.normalize("NFKC", text)
        text = re.sub(r"(\w+)-\n\s*(\w+)", r"\1\2", text)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        lines = [re.sub(r"\s+", " ", ln).strip() for ln in text.split("\n") if ln.strip()]
        return "\n".join(lines)
