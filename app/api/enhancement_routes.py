"""
API Routes for Plagiarism Platform Document Enhancement.

Implements routes for document quality analysis, paragraph-by-paragraph revisions,
applying accepted edits, version control, and report downloads.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict, List, Optional
import uuid

import docx
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from app.core.config import settings
from app.core.logging import get_logger
from app.services.extraction import TextExtractor
from app.enhancement import (
    DocumentMetricsCalculator,
    ParagraphClassifier,
    RevisionPlanner,
    ParagraphRewriter,
    DifferenceEngine,
    DocumentWriter,
    DocumentVersionManager,
    EnhancementReportCompiler,
)

logger = get_logger(__name__)
router = APIRouter(prefix="/enhancement", tags=["enhancement"])

version_manager = DocumentVersionManager()


# ── Schemas ───────────────────────────────────────────────────────────────────

class EnhancementAnalyzeRequest(BaseModel):
    file_path: str = Field(..., description="Absolute path to the uploaded document")
    filename: str = Field(..., description="Original filename of the document")


class ParagraphAnalysis(BaseModel):
    index: int
    text: str
    classification: Dict[str, Any]
    plan: Optional[Dict[str, Any]] = None


class EnhancementAnalyzeResponse(BaseModel):
    document_id: str
    filename: str
    paragraphs: List[ParagraphAnalysis]
    metrics: Dict[str, Any]


class ReviseRequest(BaseModel):
    document_id: str
    paragraph_indices: List[int] = Field(..., description="List of paragraph indices to revise")
    focus_area: str = Field("all", description="Focus of enhancement: 'all', 'originality', 'clarity', 'tone'")


class RevisionItem(BaseModel):
    index: int
    original_text: str
    revised_text: str


class ReviseResponse(BaseModel):
    document_id: str
    revisions: List[RevisionItem]


class AcceptedRevision(BaseModel):
    index: int
    revised_text: str


class ApplyRequest(BaseModel):
    document_id: str
    accepted_revisions: List[AcceptedRevision]


class ApplyResponse(BaseModel):
    document_id: str
    new_version: int
    metrics: Dict[str, Any]


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/analyze", response_model=EnhancementAnalyzeResponse)
def analyze_document_enhancement(request: EnhancementAnalyzeRequest) -> Dict[str, Any]:
    """
    Extract text, split into paragraphs, evaluate quality metrics,
    run classifiers to spot issues, and build an enhancement plan.
    """
    file_path = Path(request.file_path)
    if not file_path.exists():
        raise HTTPException(404, f"Document file not found: {request.file_path}")

    try:
        # Extract text and split by paragraphs (\n\n is standard output from TextExtractor)
        raw_text = TextExtractor.detect_and_extract(file_path)
        raw_paragraphs = [p.strip() for p in raw_text.split("\n\n") if p.strip()]
        
        if not raw_paragraphs:
            # Fallback to single line split if no \n\n found
            raw_paragraphs = [p.strip() for p in raw_text.split("\n") if p.strip()]

        if not raw_paragraphs:
            raise HTTPException(400, "Could not extract any paragraphs from the document.")

        # Compute initial metrics
        v1_metrics = DocumentMetricsCalculator.calculate_all_metrics(raw_text)

        # Classify and plan paragraph-by-paragraph
        analyzed_paragraphs: List[Dict[str, Any]] = []
        for idx, para in enumerate(raw_paragraphs):
            cls = ParagraphClassifier.classify_paragraph(para, v1_metrics)
            
            plan = None
            if cls.get("category") != "Standard":
                plan = RevisionPlanner.generate_revision_plan(para, cls)

            analyzed_paragraphs.append({
                "index": idx,
                "text": para,
                "classification": cls,
                "plan": plan
            })

        # Save to history manager
        doc_id = str(uuid.uuid4())
        version_manager.create_initial_version(
            document_id=doc_id,
            filename=request.filename,
            file_path=file_path,
            paragraphs=raw_paragraphs,
            metrics=v1_metrics
        )

        return {
            "document_id": doc_id,
            "filename": request.filename,
            "paragraphs": analyzed_paragraphs,
            "metrics": v1_metrics
        }

    except Exception as e:
        logger.error(f"Error during enhancement analysis: {e}", exc_info=True)
        raise HTTPException(500, f"Analysis failed: {str(e)}")


@router.post("/revise", response_model=ReviseResponse)
def revise_paragraphs(request: ReviseRequest) -> Dict[str, Any]:
    """
    Generate proposed paragraph edits using LLM and the revision plan strategy.
    """
    store = version_manager.get_history(request.document_id)
    if not store:
        raise HTTPException(404, f"Document history not found for ID: {request.document_id}")

    current_v = store.history[str(store.current_version)]
    paragraphs = current_v.paragraphs

    revisions: List[Dict[str, Any]] = []
    
    # Pre-calculate overall metrics for context if rewriting
    overall_metrics = current_v.metrics

    for idx in request.paragraph_indices:
        if idx < 0 or idx >= len(paragraphs):
            logger.warning(f"Index {idx} out of range for paragraph length {len(paragraphs)}")
            continue

        orig_text = paragraphs[idx]
        
        # Determine classification and plan context on the fly
        cls = ParagraphClassifier.classify_paragraph(orig_text, overall_metrics)
        plan = RevisionPlanner.generate_revision_plan(orig_text, cls)
        
        strategy = plan.get("strategy", "Enhance clarity, structure, and flow.")

        try:
            revised = ParagraphRewriter.rewrite_paragraph(
                original_text=orig_text,
                plan_strategy=strategy,
                focus_area=request.focus_area
            )
            revisions.append({
                "index": idx,
                "original_text": orig_text,
                "revised_text": revised
            })
        except Exception as e:
            logger.error(f"Failed to revise paragraph at index {idx}: {e}")
            revisions.append({
                "index": idx,
                "original_text": orig_text,
                "revised_text": orig_text  # Fallback to original text on failure
            })

    return {
        "document_id": request.document_id,
        "revisions": revisions
    }


@router.post("/apply", response_model=ApplyResponse)
def apply_revisions(request: ApplyRequest) -> Dict[str, Any]:
    """
    Accept chosen paragraph edits, rewrite files preserving styles,
    and generate a new document version.
    """
    store = version_manager.get_history(request.document_id)
    if not store:
        raise HTTPException(404, f"Document history not found for ID: {request.document_id}")

    current_v = store.history[str(store.current_version)]
    paragraphs = list(current_v.paragraphs)

    # Perform in-memory paragraph replacements
    replacements_applied: Dict[str, str] = {}
    for item in request.accepted_revisions:
        idx = item.index
        if 0 <= idx < len(paragraphs):
            orig = paragraphs[idx]
            paragraphs[idx] = item.revised_text
            replacements_applied[orig] = item.revised_text

    # Re-assemble text
    full_text = "\n\n".join(paragraphs)
    new_metrics = DocumentMetricsCalculator.calculate_all_metrics(full_text)

    # Save physical file using DocumentWriter
    active_path = Path(store.original_file_path)
    extension = active_path.suffix.lower()

    if extension == ".docx":
        # Write to DOCX in place, preserving layout and styling
        DocumentWriter.write_docx(
            original_path=Path(current_v.file_path),
            output_path=active_path,
            replacements=replacements_applied
        )
    else:
        # Default fallback to building a new styled PDF
        DocumentWriter.write_pdf(
            revised_paragraphs=paragraphs,
            output_path=active_path,
            doc_title=store.original_filename
        )

    # Record the new version inside manager
    new_v = version_manager.add_version(
        document_id=request.document_id,
        file_path=active_path,
        paragraphs=paragraphs,
        replacements=replacements_applied,
        metrics=new_metrics
    )

    return {
        "document_id": request.document_id,
        "new_version": new_v.version_num,
        "metrics": new_metrics
    }


@router.get("/report")
def get_enhancement_report(
    document_id: str = Query(...),
    version: Optional[int] = Query(None)
) -> Dict[str, Any]:
    """Retrieve comparison metrics report between v1 and a specified version."""
    store = version_manager.get_history(document_id)
    if not store:
        raise HTTPException(404, f"Document history not found for ID: {document_id}")

    v1 = store.history["1"]
    
    target_v_num = version or store.current_version
    v_target = store.history.get(str(target_v_num))
    if not v_target:
        raise HTTPException(404, f"Version {target_v_num} not found in history.")

    comparison = EnhancementReportCompiler.compile_metrics_comparison(v1, v_target)
    return {
        "document_id": document_id,
        "original_version": 1,
        "comparison_version": target_v_num,
        "comparison": comparison
    }


@router.get("/diff")
def get_document_diffs(
    document_id: str = Query(...),
    v_old: int = Query(1),
    v_new: Optional[int] = Query(None)
) -> Dict[str, Any]:
    """Get paragraph-by-paragraph differences (including inline HTML highlighted words)."""
    store = version_manager.get_history(document_id)
    if not store:
        raise HTTPException(404, f"Document history not found for ID: {document_id}")

    old_ver = store.history.get(str(v_old))
    if not old_ver:
        raise HTTPException(404, f"Source version {v_old} not found.")

    target_v_num = v_new or store.current_version
    new_ver = store.history.get(str(target_v_num))
    if not new_ver:
        raise HTTPException(404, f"Target version {target_v_num} not found.")

    diffs = DifferenceEngine.compare_documents(old_ver.paragraphs, new_ver.paragraphs)
    return {
        "document_id": document_id,
        "source_version": v_old,
        "target_version": target_v_num,
        "diffs": diffs
    }


@router.get("/download/enhanced-docx")
def download_enhanced_docx(
    document_id: str = Query(...),
    version: Optional[int] = Query(None)
) -> FileResponse:
    """Download the enhanced DOCX document of a specific version."""
    store = version_manager.get_history(document_id)
    if not store:
        raise HTTPException(404, f"Document history not found for ID: {document_id}")

    target_v_num = version or store.current_version
    v_target = store.history.get(str(target_v_num))
    if not v_target:
        raise HTTPException(404, f"Version {target_v_num} not found.")

    path = Path(v_target.file_path)
    # Ensure it's a DOCX file (or write it to DOCX)
    if path.suffix.lower() != ".docx":
        # If the original file wasn't a docx, convert/save revised paragraphs to docx
        temp_docx_path = path.parent / f"{document_id}_v{target_v_num}_export.docx"
        if not temp_docx_path.exists():
            doc = docx.Document()
            for p in v_target.paragraphs:
                doc.add_paragraph(p)
            doc.save(str(temp_docx_path))
        path = temp_docx_path

    return FileResponse(
        path=path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"enhanced_{store.original_filename}",
        headers={"Content-Disposition": f"attachment; filename=enhanced_{store.original_filename}"}
    )


@router.get("/download/enhanced-pdf")
def download_enhanced_pdf(
    document_id: str = Query(...),
    version: Optional[int] = Query(None)
) -> FileResponse:
    """Download the enhanced PDF document of a specific version."""
    store = version_manager.get_history(document_id)
    if not store:
        raise HTTPException(404, f"Document history not found for ID: {document_id}")

    target_v_num = version or store.current_version
    v_target = store.history.get(str(target_v_num))
    if not v_target:
        raise HTTPException(404, f"Version {target_v_num} not found.")

    path = Path(v_target.file_path)
    # Check if the file is a PDF, if not, generate it using DocumentWriter
    if path.suffix.lower() != ".pdf":
        temp_pdf_path = path.parent / f"{document_id}_v{target_v_num}_export.pdf"
        if not temp_pdf_path.exists():
            DocumentWriter.write_pdf(
                revised_paragraphs=v_target.paragraphs,
                output_path=temp_pdf_path,
                doc_title=store.original_filename
            )
        path = temp_pdf_path

    return FileResponse(
        path=path,
        media_type="application/pdf",
        filename=f"enhanced_{Path(store.original_filename).stem}.pdf",
        headers={"Content-Disposition": f"inline; filename=enhanced_{Path(store.original_filename).stem}.pdf"}
    )


@router.get("/download/diff-report")
def download_diff_report(
    document_id: str = Query(...),
    v_old: int = Query(1),
    v_new: Optional[int] = Query(None)
) -> FileResponse:
    """Download the PDF comparison diff report between two versions."""
    store = version_manager.get_history(document_id)
    if not store:
        raise HTTPException(404, f"Document history not found for ID: {document_id}")

    old_ver = store.history.get(str(v_old))
    if not old_ver:
        raise HTTPException(404, f"Source version {v_old} not found.")

    target_v_num = v_new or store.current_version
    new_ver = store.history.get(str(target_v_num))
    if not new_ver:
        raise HTTPException(404, f"Target version {target_v_num} not found.")

    pdf_path = version_manager.storage_dir / f"{document_id}_diff_{v_old}_to_{target_v_num}.pdf"
    
    # Compute paragraph differences
    diffs = DifferenceEngine.compare_documents(old_ver.paragraphs, new_ver.paragraphs)
    
    # Generate the Comparison PDF report
    try:
        EnhancementReportCompiler.generate_comparison_pdf(
            v1=old_ver,
            v2=new_ver,
            output_path=pdf_path,
            diff_html_paras=diffs
        )
    except Exception as e:
        logger.error(f"Failed to generate diff report PDF: {e}", exc_info=True)
        raise HTTPException(500, f"Failed to generate PDF comparison report: {str(e)}")

    return FileResponse(
        path=pdf_path,
        media_type="application/pdf",
        filename=f"comparison_report_v{v_old}_to_v{target_v_num}.pdf",
        headers={"Content-Disposition": f"attachment; filename=comparison_report_v{v_old}_to_v{target_v_num}.pdf"}
    )


@router.post("/restore")
def restore_version(
    document_id: str = Query(...),
    version: int = Query(...)
) -> Dict[str, Any]:
    """
    Restore document state to a previous version, saving a new version history entry.
    """
    try:
        restored_v = version_manager.restore_version(document_id, version)
        
        # Build paragraph analysis details
        paragraphs_info = []
        for idx, para in enumerate(restored_v.paragraphs):
            cls = ParagraphClassifier.classify_paragraph(para, restored_v.metrics)
            plan = None
            if cls.get("category") != "Standard":
                plan = RevisionPlanner.generate_revision_plan(para, cls)
            paragraphs_info.append({
                "index": idx,
                "text": para,
                "classification": cls,
                "plan": plan
            })
            
        return {
            "document_id": document_id,
            "new_version": restored_v.version_num,
            "metrics": restored_v.metrics,
            "paragraphs": paragraphs_info
        }
    except Exception as e:
        logger.error(f"Failed to restore version: {e}", exc_info=True)
        raise HTTPException(500, f"Restore failed: {str(e)}")
