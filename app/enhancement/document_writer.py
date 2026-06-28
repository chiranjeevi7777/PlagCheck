"""
Document Writer Service.

Exports enhanced documents into DOCX, PDF, HTML, and Markdown formats,
preserving original document styles, layout, and table structures.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Dict, List, Optional

import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentWriter:
    """Writes revised and enhanced text back into various target formats."""

    @staticmethod
    def write_docx(
        original_path: Path, output_path: Path, replacements: Dict[str, str]
    ) -> None:
        """
        Open original DOCX file, replace paragraph-level texts that have revisions,
        and save to output_path while preserving fonts, styles, tables, headers, and footers.
        """
        logger.info(f"Writing enhanced DOCX to {output_path}")
        if not original_path.exists():
            raise FileNotFoundError(f"Original DOCX file not found: {original_path}")

        doc = docx.Document(str(original_path))

        def clean_key(text: str) -> str:
            # Normalize whitespace for robust matching
            return re.sub(r"\s+", " ", text.strip())

        normalized_replacements = {clean_key(k): v for k, v in replacements.items()}

        def replace_in_paragraph(p: docx.text.paragraph.Paragraph) -> None:
            raw_text = clean_key(p.text)
            if not raw_text:
                return

            if raw_text in normalized_replacements:
                revised_text = normalized_replacements[raw_text]
                logger.info(f"Replacing paragraph text (length {len(raw_text)} -> {len(revised_text)})")
                
                # Replace text on the first run, empty out remaining runs to preserve styling
                if p.runs:
                    p.runs[0].text = revised_text
                    for run in p.runs[1:]:
                        run.text = ""
                else:
                    p.add_run(revised_text)

        # 1. Process root paragraphs
        for p in doc.paragraphs:
            replace_in_paragraph(p)

        # 2. Process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

        # 3. Process headers and footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    replace_in_paragraph(p)
            if section.footer:
                for p in section.footer.paragraphs:
                    replace_in_paragraph(p)

        # Ensure parent directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"DOCX document enhanced successfully: {output_path.name}")

    @staticmethod
    def write_pdf(
        revised_paragraphs: List[str], output_path: Path, doc_title: str = "Enhanced Document"
    ) -> None:
        """
        Generate a high-fidelity PDF from revised paragraphs using ReportLab.
        Detects headings and section hierarchy to apply proper styling.
        """
        logger.info(f"Writing enhanced PDF to {output_path}")
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        # Define high-end styles
        title_style = ParagraphStyle(
            'PdfTitle',
            parent=styles['Heading1'],
            fontSize=22,
            leading=26,
            textColor=colors.HexColor('#1E1E2F'),
            spaceAfter=15
        )
        heading_style = ParagraphStyle(
            'PdfHeading',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#2A2D34'),
            spaceBefore=12,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            'PdfBody',
            parent=styles['Normal'],
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=8
        )

        story = []
        story.append(Paragraph(doc_title, title_style))
        story.append(Spacer(1, 10))

        for para in revised_paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Simple heuristic to identify headings:
            # Short lines without ending punctuation or starting with section numbering
            is_heading = (
                (len(para.split()) < 12 and not para.endswith('.') and not para.endswith('?') and not para.endswith('!'))
                or re.match(r"^(?:\d+\.|\d+\.\d+|\b[A-Z\s]{4,})\b", para)
            )
            
            if is_heading:
                story.append(Paragraph(para, heading_style))
            else:
                story.append(Paragraph(para, body_style))

        doc.build(story)
        logger.info(f"PDF document generated: {output_path.name}")

    @staticmethod
    def write_html(revised_paragraphs: List[str], output_path: Path, doc_title: str = "Enhanced Document") -> None:
        """Export revised text into a premium CSS-styled HTML page."""
        logger.info(f"Writing enhanced HTML to {output_path}")
        output_path.parent.mkdir(parents=True, exist_ok=True)

        html_paragraphs = []
        for para in revised_paragraphs:
            para = para.strip()
            if not para:
                continue
            # Check heading
            is_heading = (
                (len(para.split()) < 12 and not para.endswith('.') and not para.endswith('?'))
                or re.match(r"^(?:\d+\.|\d+\.\d+|\b[A-Z\s]{4,})\b", para)
            )
            if is_heading:
                html_paragraphs.append(f"<h2>{para}</h2>")
            else:
                html_paragraphs.append(f"<p>{para}</p>")

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{doc_title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', system-ui, sans-serif;
            line-height: 1.6;
            color: #2c3e50;
            max-width: 800px;
            margin: 40px auto;
            padding: 0 20px;
            background-color: #f8fafc;
        }}
        .document {{
            background: #ffffff;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1), 0 2px 4px -1px rgba(0,0,0,0.06);
        }}
        h1 {{
            font-size: 2.25rem;
            color: #1e293b;
            margin-bottom: 20px;
            border-bottom: 2px solid #e2e8f0;
            padding-bottom: 10px;
        }}
        h2 {{
            font-size: 1.5rem;
            color: #334155;
            margin-top: 30px;
            margin-bottom: 10px;
        }}
        p {{
            margin-bottom: 16px;
            text-align: justify;
        }}
    </style>
</head>
<body>
    <div class="document">
        <h1>{doc_title}</h1>
        {"".join(html_paragraphs)}
    </div>
</body>
</html>
"""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        logger.info(f"HTML document saved: {output_path.name}")

    @staticmethod
    def write_markdown(revised_paragraphs: List[str], output_path: Path, doc_title: str = "Enhanced Document") -> None:
        """Export revised text into standard markdown."""
        logger.info(f"Writing enhanced Markdown to {output_path}")
        output_path.parent.mkdir(parents=True, exist_ok=True)

        md_lines = [f"# {doc_title}\n"]
        for para in revised_paragraphs:
            para = para.strip()
            if not para:
                continue
            is_heading = (
                (len(para.split()) < 12 and not para.endswith('.') and not para.endswith('?'))
                or re.match(r"^(?:\d+\.|\d+\.\d+|\b[A-Z\s]{4,})\b", para)
            )
            if is_heading:
                md_lines.append(f"\n## {para}\n")
            else:
                md_lines.append(f"\n{para}\n")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("".join(md_lines))
        logger.info(f"Markdown document saved: {output_path.name}")
