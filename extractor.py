import re
import unicodedata
from pathlib import Path
import docx
import fitz  # PyMuPDF
import pdfplumber
from utils import logger

class TextExtractor:
    """Handles text extraction and sanitization from PDF and DOCX files."""

    @staticmethod
    def detect_and_extract(file_path: Path) -> str:
        """Detect file type and extract text, then perform cleaning."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            logger.info(f"Extracting PDF: {file_path.name}")
            raw_text = TextExtractor.extract_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            logger.info(f"Extracting DOCX: {file_path.name}")
            raw_text = TextExtractor.extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        return TextExtractor.clean_text(raw_text)

    @staticmethod
    def extract_pdf(file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF (fitz) with fallback to pdfplumber."""
        extracted_pages = []
        try:
            doc = fitz.open(str(file_path))
            for page_num, page in enumerate(doc):
                # Page dimensions
                rect = page.rect
                height = rect.height
                width = rect.width
                
                # Heuristic margins (top and bottom 8%)
                top_margin = height * 0.08
                bottom_margin = height * 0.92
                
                blocks = page.get_text("blocks")
                page_text_parts = []
                
                for b in blocks:
                    x0, y0, x1, y1, text, block_no, block_type = b
                    # Skip image blocks
                    if block_type != 0:
                        continue
                    
                    clean_block_text = text.strip()
                    if not clean_block_text:
                        continue
                        
                    # Skip headers and footers based on coordinates
                    is_header = y0 < top_margin
                    is_footer = y1 > bottom_margin
                    
                    if is_header or is_footer:
                        # Check if it looks like a page number (only digits) or a short document label
                        if len(clean_block_text) < 15 or re.match(r'^\d+$', clean_block_text) or re.match(r'^page\s+\d+\s*(of\s*\d+)?$', clean_block_text, re.IGNORECASE):
                            continue
                            
                    page_text_parts.append(clean_block_text)
                
                extracted_pages.append("\n".join(page_text_parts))
            
            combined_text = "\n\n".join(extracted_pages)
            if combined_text.strip():
                return combined_text
                
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed for {file_path.name}, falling back to pdfplumber: {e}")

        # Fallback to pdfplumber
        try:
            with pdfplumber.open(str(file_path)) as pdf:
                extracted_pages = []
                for page in pdf.pages:
                    # Basic extraction
                    text = page.extract_text()
                    if text:
                        # Basic line-by-line filtering for page numbers/headers
                        lines = text.split("\n")
                        filtered_lines = []
                        for line in lines:
                            line_stripped = line.strip()
                            if not line_stripped:
                                continue
                            # Exclude single numbers or common page headers
                            if re.match(r'^\d+$', line_stripped) or re.match(r'^page\s+\d+\s*(of\s*\d+)?$', line_stripped, re.IGNORECASE):
                                continue
                            filtered_lines.append(line_stripped)
                        extracted_pages.append("\n".join(filtered_lines))
                return "\n\n".join(extracted_pages)
        except Exception as e:
            logger.error(f"pdfplumber extraction failed as well: {e}")
            raise RuntimeError(f"Could not extract text from PDF: {e}")

    @staticmethod
    def extract_docx(file_path: Path) -> str:
        """Extract text from DOCX including paragraphs and tables."""
        try:
            doc = docx.Document(str(file_path))
            text_parts = []
            
            # Extract paragraphs
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    text_parts.append(p_text)
                    
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        text_parts.append(" | ".join(row_cells))
                        
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"python-docx extraction failed: {e}")
            raise RuntimeError(f"Could not extract text from DOCX: {e}")

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize the extracted text according to specifications."""
        if not text:
            return ""
            
        # 1. Normalize Unicode (NFKC handles ligatures and compatibility characters)
        text = unicodedata.normalize("NFKC", text)
        
        # 2. Fix hyphenated word breaks at line ends (e.g., "plagi-\narism" -> "plagiarism")
        text = re.sub(r'(\w+)-\n\s*(\w+)', r'\1\2', text)
        
        # 3. Replace vertical whitespaces or line-breaks with clean newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 4. Remove multiple spaces and collapse empty lines
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            # Replace multiple spaces/tabs with single space
            line_clean = re.sub(r'\s+', ' ', line).strip()
            if line_clean:
                cleaned_lines.append(line_clean)
                
        # Return combined clean lines separated by space (or newline if paragraph structure is preferred)
        # To maintain paragraph boundaries, let's keep them separated by newlines
        return "\n".join(cleaned_lines)
